from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
)
from reportlab.lib.styles import getSampleStyleSheet


GENERATED_DIR = Path("generated")
GENERATED_DIR.mkdir(exist_ok=True)


def generate_docx(
    meeting_id: int,
    title: str,
    minutes: str,
) -> str:

    path = GENERATED_DIR / f"meeting-{meeting_id}.docx"

    document = Document()

    document.add_heading(title, level=1)
    document.add_heading("Meeting Minutes", level=2)

    for paragraph in minutes.split("\n\n"):
        document.add_paragraph(paragraph)

    document.save(path)

    return str(path)


def generate_pdf(
    meeting_id: int,
    title: str,
    minutes: str,
) -> str:

    path = GENERATED_DIR / f"meeting-{meeting_id}.pdf"

    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
    )

    styles = getSampleStyleSheet()

    story = [
        Paragraph(title, styles["Title"]),
        Spacer(1, 20),
    ]

    for paragraph in minutes.split("\n\n"):
        story.append(
            Paragraph(
                paragraph.replace("\n", "<br/>"),
                styles["BodyText"],
            )
        )

        story.append(Spacer(1, 10))

    document.build(story)

    return str(path)